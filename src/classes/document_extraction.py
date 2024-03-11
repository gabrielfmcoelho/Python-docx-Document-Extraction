#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Oct 22 20:59:46 2019
Modified on Fri Mar 08 15:37:00 2024

@author: karthick
@modified by: Gabriel Coelho
"""
# ------------------------------
## >> EXTERNAL DEPENDENCIES
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document as Document_compose
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
from typing import Dict
import io
import csv
from icecream import ic
import base64
import os

# ------------------------------
## >> CUSTOM DEPENDENCIES
from document_output import DocumentOutput


class DocumentExtraction:
    # ------------------------------
    ## >> ATTRIBUTES
    ### * The raw_document will store the original document object
    raw_document = None

    ### * The df_document_content dataframe will store all the references to the content in document order by paragraphs.
    df_document_content = pd.DataFrame(columns=['document_name', 'paragraph_content', 'content_reference_id', 'style', 'style_extracted', 'highlighted_content'])
    
    ### * The df_document_resources will consist of base64 encoded image data of all the images in the document and the corresponding image id
    df_document_resources = pd.DataFrame(columns=['resource_index', 'image_rID', 'image_filename', 'image_base64_string', 'resource_type', 'text_content'])

    #xml_list=[]

    # ------------------------------
    ## >> CONSTRUCTOR
    ### [1] The constructor will load the document object from the input path
    def __init__(self, document_path: str) -> None:
        try:
            self.raw_document = Document(document_path)
            self.document_name = os.path.basename(document_path)
            self.resource_counter = 0
            self.table_counter = 0
        except FileNotFoundError:
            raise ValueError('The document was not found. Please, check the path and try again.')

    # ------------------------------
    ## >> PRIVATE METHODS
    ### [x] This function extracts the tables and paragraphs from the document object
    def __iter_block_items(self, parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        #### Access element directly (to access body, a child of _Cell)
        if isinstance(parent, doctwo):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc 
        else:
            raise ValueError("something's not right, thats not a document or a cell")
        #### Iterate through every child of the element extracted
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    ### [x] This function extracts the table from the document object as a dataframe90
    def __read_docx_tables(self, tab_id=None, **kwargs):
        """
        parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

        Parameters:
            filename:   file name of a Word Document

            tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                        When [None] - return a list of DataFrames (parse all tables)

            kwargs:     arguments to pass to `pd.read_csv()` function

        Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
        """
        def __read_docx_tab(tab, **kwargs):
            vf = io.StringIO()
            writer = csv.writer(vf)
            for row in tab.rows:
                writer.writerow(cell.text for cell in row.cells)
            vf.seek(0)
            # read vf to a dataframe, the header must be specified to avoid the first row being used as header
            df = pd.read_csv(vf, **kwargs, header=None)
            if len(df.columns) == 1:
                column = df.columns.to_list()
                column = column[0]
                return str(df.iloc[column, 0])
            return df

        if tab_id is None:
            return [__read_docx_tab(tab, **kwargs) for tab in self.raw_document.tables]
        else:
            try:
                return __read_docx_tab(self.raw_document.tables[tab_id], **kwargs)
            except IndexError:
                print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                raise

    def __process_image_xml(self, xml_str) -> str:
        #ic(xml_str)
        my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xml_str), events=['start-ns'])])
        root = ET.fromstring(xml_str) 
        #self.xml_list.append(xml_str) # NOTE: SAVING THE XML STRING FOR FUTURE USE BUT NOT USED IN THIS CODE
        for pic in root.findall('.//pic:pic', my_namespaces):
            cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
            name_attr = cNvPr_elem.get("name")
            blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
            embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(self.resource_counter))
            document_part = self.raw_document.part
            image_part = document_part.related_parts[embed_attr]
            image_base64 = base64.b64encode(image_part._blob)
            image_base64 = image_base64.decode()                            
            self.df_document_resources = self.df_document_resources._append({
                'resource_index':self.resource_counter,
                'image_rID':embed_attr,
                'image_filename':name_attr,
                'image_base64_string':image_base64,
                'resource_type': 'image'
                }, sort=False, ignore_index=True)
        counter = self.resource_counter
        self.__update_resource_counter()
        return None, counter, 'Resource'

    def __process_text_block(self, block: Paragraph) -> None:            
        #ic(block.style)
        style = str(block.style.name)
        #ic(style)
        appendtxt = str(block.text)
        appendtxt = appendtxt.replace("\n","")
        appendtxt = appendtxt.replace("\r","")
        content_reference = 'Novalue'
        #ic(appendtxt)
        paragraph_split = appendtxt.lower().split() #NOTE: NO USE !!!

        styles = {}
        styles['alignment'] = block.alignment
        try:
            styles['spacing_before'] = block.paragraph_format.space_before.pt
        except:
            styles['spacing_before'] = None
        try:
            styles['spacing_after'] = block.paragraph_format.space_after.pt
        except:
            styles['spacing_after'] = None
        styles['runs'] = []
        self.has_highlighted = False

        for run in block.runs:
            def get_run_styles(run):
                styles = {}
                styles['font_name'] = run.font.name
                styles['font_size'] = run.font.size
                styles['bold'] = run.bold
                styles['italic'] = run.italic
                styles['underline'] = run.underline
                styles['color'] = run.font.color.rgb
                styles['highlight'] = run.font.highlight_color
                if run.font.highlight_color:
                    self.has_highlighted = True
                if styles:
                    return styles
                return None
            styles['runs'] = styles['runs'] + [get_run_styles(run)]
            if 'pic:pic' in str(run.element.xml): # DEAL WITH IMAGES >> Check if pic is there in the xml of the element. If yes, then extract the image data
                appendtxt, content_reference, style = self.__process_image_xml(str(run.element.xml))

        #ic(styles)
        #ic(self.has_highlighted)
        if (appendtxt != '' and appendtxt != ' ' and appendtxt != '\n' and appendtxt != '\r'):
            self.__append_content(appendtxt, content_reference, style, styles, self.has_highlighted)

    def __process_table_block(self, block: Table) -> None:
        dfs = self.__read_docx_tables(tab_id=self.table_counter)
        self.df_document_resources = self.df_document_resources._append({
            'resource_index':self.resource_counter,
            'text_content': dfs,
            'resource_type': 'table'
            }, sort=False, ignore_index=True)
        self.table_counter += 1
        self.__append_content(None, self.resource_counter, 'Resource')
        self.__update_resource_counter()

    def __append_content(self, content: str, content_reference_id: int, style: str, style_detailed: dict = None, has_highlighted: bool = None) -> None:
        self.df_document_content = self.df_document_content._append({
            'document_name': self.document_name,
            'paragraph_content': content,
            'content_reference_id': content_reference_id,
            'style': style,
            'style_extracted': style_detailed,
            'highlighted_content': has_highlighted
            }, sort=False, ignore_index=True)
        
    def __update_resource_counter(self) -> None:
        self.resource_counter += 1
    
    ### [x] This function receives the created document objects and stores them in the document_collection as a single data structure
    def __update_document_collection(self) -> None:
        self.document_collection = {
            'content': self.df_document_content,
            'resources': self.df_document_resources
            }
    
    # ------------------------------
    ## >> MAIN METHODS
    ### [2] This function extracts the content from the document object
    def extract(self):
        processing_methods = {
            'text': self.__process_text_block,
            'table': self.__process_table_block
        }

        for block in self.__iter_block_items(self.raw_document):
            #ic(block) # i cannot see the output of this print statement only a object reference
            # like docx.text.paragraph.Paragraph object at 0x7f8e3c3e3e50 lets see if it changes
            for method in processing_methods:
                if method in str(block):
                    #ic(method)
                    processing_methods[method](block)
                    #ic(self.df_document_content)
                    #ic(self.df_document_resources)
                    #ic(self.xml_list)
                    break
        
        #ic(self.df_document_content)
        #ic(self.df_document_resources)
        #ic(self.xml_list)
        self.__update_document_collection()
        return self.document_collection

    ### [3] This function exports the document content and references to files
    def export(self, output_format: DocumentOutput, document_output_path: str = None, document_collection: dict = None) -> None:
        def __build_document_output_path(document_output_path: str, document_name: str, document_output_type: str) -> str:
            if document_output_path is None:
                document_output_path = ''
            raw_document_name = self.document_name.replace(".docx", '_')
            return document_output_path + raw_document_name + document_name + document_output_type
            
        if document_collection is None:
            try:
                document_collection = self.document_collection
            except AttributeError:
                raise ValueError('The document_collection is not defined. Please, extract the document first.')
        
        for document_name, document in document_collection.items():
            output_path = __build_document_output_path(document_output_path, document_name, output_format.value['filetype'])
            ic('saving document at path:', output_path)
            output_format.value['function'](document, output_path, **output_format.value['args'])